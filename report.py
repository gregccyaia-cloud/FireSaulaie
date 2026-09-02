from datetime import datetime,timezone
import os,time
from docx import Document
from docx.shared import Cm,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def fig(d,p,c,w=15.5):
 d.add_picture(str(p),width=Cm(w));d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER;q=d.add_paragraph(c);q.style='Caption';q.alignment=WD_ALIGN_PARAGRAPH.CENTER
def table(d,df,fs=8):
 h=list(df.columns);t=d.add_table(rows=1,cols=len(h));t.style='Table Grid'
 for i,x in enumerate(h):t.rows[0].cells[i].text=str(x)
 for _,r in df.iterrows():
  c=t.add_row().cells
  for i,x in enumerate(r):c[i].text=str(x)
 for row in t.rows:
  for c in row.cells:
   for p in c.paragraphs:
    for x in p.runs:x.font.size=Pt(fs)
def toc(p):
 r=p.add_run();b=OxmlElement('w:fldChar');b.set(qn('w:fldCharType'),'begin');i=OxmlElement('w:instrText');i.set(qn('xml:space'),'preserve');i.text=' TOC \\o "1-3" \\h \\z ';s=OxmlElement('w:fldChar');s.set(qn('w:fldCharType'),'separate');t=OxmlElement('w:t');t.text='Mettre à jour dans Word : Ctrl+A, F9.';e=OxmlElement('w:fldChar');e.set(qn('w:fldCharType'),'end')
 for x in (b,i,s,t,e):r._r.append(x)
def build(ctx,out):
 d=Document();sec=d.sections[0];sec.top_margin=sec.bottom_margin=Cm(1.6);sec.left_margin=sec.right_margin=Cm(2);d.styles['Normal'].font.name='Arial';d.styles['Normal'].font.size=Pt(9.5)
 for n in ('Title','Heading 1','Heading 2'):d.styles[n].font.color.rgb=RGBColor(31,78,120)
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('ANALYSE THERMIQUE D’UN INCENDIE SOUS LA PASSERELLE');r.bold=True;r.font.size=Pt(18);p=d.add_paragraph('Passerelle Gerland - La Saulaie');p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p=d.add_paragraph('Rapport du '+datetime.now().strftime('%d/%m/%Y'));p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 d.add_page_break();d.add_heading('Sommaire',1);toc(d.add_paragraph());d.add_page_break()
 d.add_heading('1. Objet et présentation du projet',1);d.add_paragraph('L’étude évalue l’évolution temporelle des températures des suspentes secondaires et des câbles principaux sous l’action d’un incendie sur la M7. Les calculs, l’évolution des propriétés thermiques et l’intégration pas à pas sont développés sous Python.')
 for p,c,w in ctx['assets_intro']:fig(d,p,c,w)
 d.add_heading('2. Références',1);d.add_paragraph('• NF EN 1991-1-2 : actions thermiques.\n• NF EN 1993-1-2 : comportement au feu de l’acier.\n• Cerema, Résistance à l’incendie des ponts routiers, 2018.\n• Note SAU_AVP_NTE_069_A_FeuM7.')
 d.add_heading('3. Choix de la courbe de feu',1);d.add_paragraph(ctx['choice']);fig(d,ctx['cerema'],'Figure 6 - Extrait du guide Cerema : choix des courbes.',15);d.add_paragraph(ctx['moa']);fig(d,ctx['fire'],'Figure 7 - Courbes nominales température-temps.')
 d.add_heading('4. Géométrie simplifiée et positions de feu',1);table(d,ctx['geometry'],8);fig(d,ctx['intrados'],'Figure 8 - Approximation de l’intrados par deux segments de droite et points géométriques de calage.',14);d.add_paragraph('F1 correspond à la zone ouest, F2 à l’axe de la M7 et F3 à la zone est.');fig(d,ctx['geom'],'Figure 9 - Trois coupes géométriques de référence et positions F1, F2 et F3.')
 d.add_heading('5. Choix des paramètres et hypothèses',1);table(d,ctx['params'],8);d.add_paragraph(ctx['cp'])
 d.add_heading('6. Méthode de calcul',1);d.add_paragraph('Pour chaque courbe, position et longueur de foyer, Python calcule la température des gaz, le facteur de forme, la chaleur spécifique dépendante de la température, les flux convectif et radiatif, puis l’incrément de température. L’intégration est répétée par pas de 5 s. L’annexe A illustre les étapes et le dernier pas avant 30 min.')
 d.add_heading('7. Résultats',1);fig(d,ctx['f10'],'Figure 10 - Échauffement - Suspente secondaire - foyer 20 m.');fig(d,ctx['f11'],'Figure 11 - Échauffement - Câble principal clos - foyer 20 m.');d.add_paragraph('Les valeurs ci-dessous correspondent aux enveloppes maximales de tous les cas calculés pour chaque famille d’éléments.');table(d,ctx['retained'],8)
 d.add_heading('8. Impact structurel',1);d.add_paragraph('Les vérifications structurelles devront être menées aux temps 15, 30, 60, 90 et 120 min, à partir des températures enveloppes du § 7 et des efforts axiaux quasi permanents à fournir pour les câbles principaux et les suspentes secondaires. Pour chaque élément et chaque temps, il conviendra de :')
 for x in ['déterminer les coefficients de réduction de la limite d’élasticité et du module d’Young à la température retenue ;','calculer la résistance axiale à chaud de la section, en tenant compte de la section efficace et du coefficient partiel applicable en situation d’incendie ;','établir le ratio d’utilisation ηfi = Nfi,Ed / Nfi,Rd,θ et vérifier qu’il reste inférieur ou égal à 1,00 ;','évaluer l’allongement thermique, la perte de rigidité et la variation de tension ;','examiner les redistributions d’efforts entre suspentes et câbles principaux ainsi que, si nécessaire, les scénarios de perte locale d’une suspente.']:d.add_paragraph(x,style='List Bullet')
 d.add_paragraph('Les efforts axiaux considérés sont les efforts quasi permanents, conformément à l’hypothèse d’action accidentelle principale constituée par le feu retenue pour la poursuite de l’analyse.')
 d.add_heading('9. Portée et limites',1);d.add_paragraph('L’approche du présent rapport constitue une évaluation thermique enveloppe. Le choix d’un facteur de forme prudent et l’absence de crédit explicite pour les masquages placent le calcul du côté de la sécurité pour les configurations considérées. Un affinage géométrique pourra ultérieurement mieux différencier les positions et quantifier les effets favorables éventuels.')
 d.add_paragraph('Pour les suspentes, l’approche considère actuellement une exposition directe sans protection. Or, leur naissance est partiellement située derrière le bec du tablier, qui peut produire un effet de masque vis-à-vis d’une partie du rayonnement. Cette hypothèse est donc prudente et pourra être affinée par un calcul de visibilité et de facteur de forme. Elle ne permet toutefois pas d’écarter une exposition significative, puisqu’un foyer étendu peut également rayonner latéralement vers les suspentes.')
 d.add_paragraph('Le scénario ne couvre pas un feu de citerne, car le présent rapport n’étudie pas la courbe HC ; voir le § 3 « Choix de la courbe de feu ».')
 d.add_page_break();d.add_heading('Annexe A - Calcul détaillé du cas critique à 30 min',1);table(d,ctx['steps'],8);fig(d,ctx['integration'],'Figure A.1 - Intégration temporelle du cas critique.');table(d,ctx['worst'],8)
 d.add_page_break();d.add_heading('Annexe B - Courbes de tous les cas étudiés',1);d.add_paragraph('Les figures présentent les neuf cas F1/F2/F3 et L = 10/15/20 m, le domaine min.-max. et l’enveloppe maximale.')
 for p,c in ctx['annex']:fig(d,p,c,15)
 now=datetime.now(timezone.utc);d.core_properties.created=now;d.core_properties.modified=now;u=OxmlElement('w:updateFields');u.set(qn('w:val'),'true');d.settings._element.append(u);out.parent.mkdir(parents=True,exist_ok=True);d.save(out);os.utime(out,(time.time(),time.time()));return out
