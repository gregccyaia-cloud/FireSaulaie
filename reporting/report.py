from docx import Document
from docx.shared import Cm,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def shade(cell,fill):
 tcPr=cell._tc.get_or_add_tcPr();shd=OxmlElement('w:shd');shd.set(qn('w:fill'),fill);tcPr.append(shd)
def table(doc,headers,rows,widths=None):
 t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for i,h in enumerate(headers):t.rows[0].cells[i].text=str(h);shade(t.rows[0].cells[i],'1F4E78')
 for c in t.rows[0].cells:
  c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
  for run in c.paragraphs[0].runs:run.font.bold=True;run.font.color.rgb=RGBColor(255,255,255)
 for row in rows:
  cells=t.add_row().cells
  for i,v in enumerate(row):cells[i].text=str(v);cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
 return t

def add_fig(doc,path,caption):
 if path.exists():
  p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(path),width=Cm(15.8))
  c=doc.add_paragraph(caption);c.alignment=WD_ALIGN_PARAGRAPH.CENTER;c.style='Caption'

def create_report(ctx,output):
 output.parent.mkdir(parents=True,exist_ok=True);d=Document();sec=d.sections[0];sec.top_margin=Cm(1.8);sec.bottom_margin=Cm(1.8);sec.left_margin=Cm(2);sec.right_margin=Cm(2)
 styles=d.styles;styles['Normal'].font.name='Arial';styles['Normal'].font.size=Pt(9.5)
 for n in ('Title','Heading 1','Heading 2','Heading 3'):styles[n].font.name='Arial';styles[n].font.color.rgb=RGBColor(31,78,120)
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('ANALYSE THERMIQUE INCENDIE SOUS PASSERELLE');r.bold=True;r.font.size=Pt(20);r.font.color.rgb=RGBColor(31,78,120)
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run('Modele Python V1.2 - Rapport consolide de validation').bold=True
 d.add_paragraph('Passerelle Gerland - La Saulaie | Franchissement de la M7',style='Subtitle');d.add_paragraph('Document de travail. Les hypotheses spatiales V1.2 restent une approche de depistage et ne constituent pas un modele de feu localise normatif.')
 d.add_page_break()
 d.add_heading('1. Objet et perimetre',1);d.add_paragraph('L etude calcule l evolution temporelle de la temperature de l intrados metallique, des suspentes secondaires et du cable principal pour deux courbes nominales de feu, trois positions de camion et trois hauteurs d enveloppe de flamme. Les calculs thermiques sont dissocies des futures verifications structurelles.')
 d.add_heading('2. Documents et references',1)
 table(d,['Reference','Utilisation'],[['Guide CEREMA 2018 - Resistance a l incendie des ponts routiers','Demarche, feu exterieur et transfert thermique'],['NF EN 1991-1-2','Actions thermiques, convection et rayonnement'],['NF EN 1993-1-2','Proprietes thermiques de l acier et facteurs de reduction'],['Note projet 069','Analyse initiale du risque incendie']])
 d.add_heading('3. Donnees projet et geometrie',1)
 table(d,['Donnee','Valeur'],ctx['project_rows']);add_fig(d,ctx['geometry_png'],'Figure 1 - Geometrie longitudinale simplifiee.')
 d.add_paragraph('Le profil transversal de l intrados est defini par deux segments symetriques : remontee de 0,19 m a y = 1,91 m puis de 0,61 m au bord y = 3,70 m. La naissance inferieure des suspentes est prise a 0,82 m au-dessus du point bas de l intrados.')
 d.add_heading('4. Scenarios de calcul',1)
 d.add_paragraph('Deux calculs independants sont menes. L enveloppe n est constituee qu apres obtention des resultats de chaque courbe et de chaque position de camion.')
 table(d,['Scenario','Expression','Domaine temporel'],ctx['fire_rows']);add_fig(d,ctx['fire_png'],'Figure 2 - Courbes de temperature nominale des gaz.')
 d.add_heading('5. Modele thermique',1)
 d.add_heading('5.1 Bilan energetique',2);d.add_paragraph('Chaque element acier est traite par un modele capacitif a temperature uniforme. Sur un pas de temps Delta t, la variation de temperature resulte du flux net convectif et radiatif rapporte a la masse volumique, a la chaleur specifique et au facteur de massiveté A_m/V.')
 d.add_paragraph('Delta theta_a = (A_m/V) x (q_conv + q_rad) x Delta t / (rho_a x c_a(theta_a))')
 d.add_paragraph('q_conv = f_conv x alpha_c x (theta_g - theta_a)')
 d.add_paragraph('q_rad = f_rad x epsilon_m x epsilon_f x sigma x [(theta_g + 273,15)^4 - (theta_a + 273,15)^4]')
 table(d,['Parametre','Valeur V1.2','Statut'],ctx['thermal_rows'])
 d.add_heading('5.2 Facteurs de massiveté',2);d.add_paragraph('Pour un element circulaire pleinement expose : A_m/V = 4/D. Les valeurs retenues sont 95,24 m-1 pour la suspente de 42 mm et 30,30 m-1 pour le cable de 132 mm. La valeur de l intrados est provisoire et devra etre remplacee par la geometrie thermique du caisson.')
 d.add_heading('5.3 Representation spatiale de la flamme',2);d.add_paragraph('Les hauteurs 10, 12 et 15 m constituent une analyse de sensibilite geometrique. Le modele applique des coefficients d exposition directs ou masques selon la position de l element par rapport a cette enveloppe. Cette representation n est pas l annexe C de l EN 1991-1-2. Une etape ulterieure devra employer Q(t), le diametre equivalent D, la longueur de flamme et la distance horizontale r pour un feu localise.')
 d.add_heading('6. Etapes de calcul et valeurs intermediaires',1)
 table(d,['Etape','Operation','Sortie controlee'],ctx['steps_rows'])
 d.add_heading('7. Validation numerique',1);d.add_paragraph(ctx['validation_text']);add_fig(d,ctx['conv_png'],'Figure 3 - Comparaison des pas de temps 5 s et 2,5 s sur le cas de controle.')
 d.add_heading('8. Resultats synthetiques',1);table(d,list(ctx['summary'].columns),ctx['summary'].values.tolist())
 d.add_heading('9. Courbes representatives',1)
 for p,cap in ctx['case_figures']:add_fig(d,p,cap)
 d.add_heading('10. Conclusions et limites de validation',1)
 for txt in ctx['conclusions']:d.add_paragraph(txt,style='List Bullet')
 d.add_heading('11. Donnees requises pour la verification structurelle',1)
 for txt in ['Efforts de traction de reference dans les suspentes et le cable principal.','Nuances et caracteristiques mecaniques exactes des aciers.','Longueur chauffee et conditions d ancrage ou de redistribution.','Criteres de dommage acceptables pendant et apres incendie.']:d.add_paragraph(txt,style='List Bullet')
 d.save(str(output));return output
