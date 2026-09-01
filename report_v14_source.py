from pathlib import Path
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm,Pt,Mm

def r(x): o=OxmlElement('m:r'); t=OxmlElement('m:t'); t.text=x; o.append(t); return o
def sub(b,s): o=OxmlElement('m:sSub'); e=OxmlElement('m:e');e.append(r(b));q=OxmlElement('m:sub');q.append(r(s));o.extend([e,q]);return o
def sup(b,s): o=OxmlElement('m:sSup');e=OxmlElement('m:e');e.append(r(b));q=OxmlElement('m:sup');q.append(r(s));o.extend([e,q]);return o
def frac(n,d): o=OxmlElement('m:f');a=OxmlElement('m:num');b=OxmlElement('m:den');[a.append(x) for x in n];[b.append(x) for x in d];o.extend([a,b]);return o
def eq(doc,els): p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;mp=OxmlElement('m:oMathPara');m=OxmlElement('m:oMath');[m.append(x) for x in els];mp.append(m);p._p.append(mp)
def q(s): return sub('q',s)
def th(s): return sub('θ',s)
def image(doc,path,caption,w=16):
    path=Path(path)
    if not path.exists(): return 0
    doc.add_picture(str(path),width=Cm(w));doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph(caption);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.runs[0].italic=True;return 1
def table(doc,rows):
    if not rows:return
    hs=list(rows[0]);t=doc.add_table(rows=1,cols=len(hs));t.style='Table Grid'
    for i,h in enumerate(hs):t.rows[0].cells[i].text=h;t.rows[0].cells[i].paragraphs[0].runs[0].bold=True
    for row in rows:
        c=t.add_row().cells
        for i,h in enumerate(hs):c[i].text=str(row[h])
def f(v,n=2):return f'{v:.{n}f}'.replace('.',',')
def equations(doc):
    doc.add_heading('Équations utilisées',2)
    doc.add_paragraph('Courbe CN - ISO 834, NF EN 1991-1-2, § 3.2.1 :');eq(doc,[th('g'),r('(t) = 20 + 345 log'),sub('','10'),r('(8t + 1)')])
    doc.add_paragraph('Courbe de feu extérieur, NF EN 1991-1-2, § 3.2.2 :');eq(doc,[th('g'),r('(t) = 660[1 - 0,687'),sup('e','-0,32t'),r(' - 0,313'),sup('e','-3,8t'),r('] + 20')])
    doc.add_paragraph('Flux convectif net, NF EN 1991-1-2, § 3.1 :');eq(doc,[q('conv'),r(' = α'),sub('','c'),r('('),th('g'),r(' - '),th('a'),r(')')])
    doc.add_paragraph('Flux radiatif net, NF EN 1991-1-2, § 3.1 :');eq(doc,[q('rad'),r(' = Φ ε'),sub('','m'),r(' ε'),sub('','f'),r(' σ [('),th('g'),r(' + 273,15)'),sup('','4'),r(' - ('),th('a'),r(' + 273,15)'),sup('','4'),r(']')])
    doc.add_paragraph('Facteur de section :');eq(doc,[frac([sub('A','m')],[r('V')]),r(' = η'),sub('','exp'),r(' '),frac([r('4')],[r('D')])])
    doc.add_paragraph("Incrément de température, NF EN 1993-1-2, § 4.2.5.1 :");eq(doc,[r('Δ'),th('a,t'),r(' = '),frac([sub('k','sh'),r(' '),frac([sub('A','m')],[r('V')]),r(' '),q('net'),r(' Δt')],[sub('ρ','a'),r(' '),sub('c','a'),r('('),th('a'),r(')')])])
def build(path,figures,summary,case,assets,sections):
    d=Document();s=d.sections[0];s.top_margin=s.bottom_margin=Mm(18);s.left_margin=s.right_margin=Mm(20);d.styles['Normal'].font.name='Aptos';d.styles['Normal'].font.size=Pt(10.5)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(55);x=p.add_run("ANALYSE THERMIQUE D'UN INCENDIE\nSOUS LA PASSERELLE");x.bold=True;x.font.size=Pt(23)
    p=d.add_paragraph('Passerelle Gerland - La Saulaie\nVersion V1.4 complète');p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=d.add_paragraph('Rapport généré le '+datetime.now().strftime('%d/%m/%Y à %H:%M'));p.alignment=WD_ALIGN_PARAGRAPH.CENTER;d.add_page_break()
    d.add_heading("1. Objet et présentation du projet",1);d.add_paragraph("L'étude évalue l'évolution de la température de l'intrados, des suspentes et des câbles principaux lors d'un incendie sur la M7. Les courbes ISO 834 et feu extérieur sont calculées séparément, puis comparées.")
    ins=0;ins+=image(d,assets/'ZZ extrait vue en plan generale.png','Figure 1 - Vue en plan générale.');ins+=image(d,assets/'ZZ extrait plan suspension.png','Figure 2 - Élévation et plan du système de suspension.');ins+=image(d,assets/'Zz extrait coupe long generale.png','Figure 3 - Coupe longitudinale générale.')
    d.add_heading('2. Références',1)
    for z in ['NF EN 1991-1-2, § 3.1, § 3.2.1 et § 3.2.2.','NF EN 1993-1-2, § 3.4.1 et § 4.2.5.1.','Cerema, Résistance à l\'incendie des ponts routiers, guide 2018, p. 14/156.','Note SAU_AVP_NTE_069_A_FeuM7.']:d.add_paragraph(z,style='List Bullet')
    d.add_heading('3. Géométrie simplifiée',1);d.add_paragraph("Trois coupes successives sont étudiées le long de l'axe de la passerelle au franchissement de la M7 : bord ouest, axe M7 et bord est. Les hauteurs ne correspondent donc pas à trois points d'une même coupe transversale.")
    rows=[{'Coupe':s.code,'Position':s.label,'Intrados (m)':f(s.intrados_m,3),'Naissance suspente (m)':f(s.intrados_m+.82,3),'Axe câble (m)':f(s.cable_m,3)} for s in sections];table(d,rows)
    ins+=image(d,assets/'ZZ_coupe longit partielle zone M7.png','Figure 4 - Hauteurs sur les trois coupes.');ins+=image(d,assets/'ZZ_vue en plan cotée zone M7.png','Figure 5 - Largeurs de la M7.');ins+=image(d,assets/'ZZ_demie coupe tablier et approx intrados.png',"Figure 6 - Demi-coupe et approximation de l'intrados.");ins+=image(d,assets/'ZZ coupe transv.png','Figure 7 - Coupe transversale type.')
    d.add_heading('4. Hypothèses et modèle thermique',1);equations(d)
    d.add_heading('5. Résultats',1);table(d,summary)
    for title,pth in figures:ins+=image(d,pth,title)
    d.add_heading('6. Limites',1);d.add_paragraph("Le facteur de configuration Φ vaut 1,0. Les longueurs de foyer sont paramétrées, mais elles ne modifient pas encore le flux tant qu'un facteur de forme géométrique n'est pas introduit.")
    d.add_page_break();d.add_heading('Annexe A - Calcul détaillé du cas critique à 30 min',1)
    d.add_paragraph(f"Cas retenu : {case['fire']} - {case['position']} - coupe {case['section']} - {case['element']} - foyer {f(case['fire_length'],1)} m.")
    table(d,[{'Grandeur':k,'Valeur':v} for k,v in case['values']])
    d.add_heading('A.1 Calcul des flux',2);eq(d,[q('conv'),r(' = α'),sub('','c'),r('('),th('g'),r(' - '),th('a'),r(') = '+f(case['qc'],1)+' W/m²')]);eq(d,[q('rad'),r(' = Φ ε'),sub('','m'),r(' ε'),sub('','f'),r(' σ [('),th('g'),r(' + 273,15)'),sup('','4'),r(' - ('),th('a'),r(' + 273,15)'),sup('','4'),r('] = '+f(case['qr'],1)+' W/m²')]);eq(d,[q('net'),r(' = '),q('conv'),r(' + '),q('rad'),r(' = '+f(case['qn'],1)+' W/m²')])
    d.add_heading('A.2 Incrément de température',2);eq(d,[r('Δ'),th('a,t'),r(' = '),frac([sub('k','sh'),r(' '),frac([sub('A','m')],[r('V')]),r(' '),q('net'),r(' Δt')],[sub('ρ','a'),r(' '),sub('c','a'),r('('),th('a'),r(')')]),r(' = '+f(case['dta'],4)+' °C')]);d.add_paragraph(f"La température passe de {f(case['ta0'],2)} °C à {f(case['ta1'],2)} °C pendant le dernier pas de 5 s. La température à 30 min résulte de l'intégration de tous les pas depuis t = 0.")
    path=Path(path);path.parent.mkdir(parents=True,exist_ok=True);d.save(path);return {'path':path.resolve(),'figures':ins}
